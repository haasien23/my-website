from __future__ import annotations

import re
from pathlib import Path

import matplotlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

matplotlib.use("Agg")
import matplotlib.pyplot as plt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "pdf" / "毕业论文初稿_汇总版.md"
if not SOURCE.exists():
    SOURCE = sorted((ROOT / "output" / "pdf").glob("*.md"))[-1]

OUTPUT = ROOT / "output" / "doc" / "毕业论文初稿_汇总版_公式终版_v2.docx"
EQUATION_DIR = ROOT / "tmp" / "docs" / "equations_final"

EQUATION_RE = re.compile(r"^\(([^)]+)\)\s*(.+)$")

plt.rcParams["mathtext.fontset"] = "stix"
plt.rcParams["font.family"] = "STIXGeneral"


EQUATION_OVERRIDES: dict[str, tuple[str, str]] = {
    "6-1": ("latex", r"\mathcal{E}(t)=\left(\Omega,\, O,\, U(t),\, T(t),\, W\right)"),
    "6-2": ("piecewise", "chi"),
    "6-3": ("piecewise", "visibility"),
    "6-24": ("latex", r"\rho_i(t)=\operatorname{clip}\left(\frac{N_i^{\mathrm{nbr}}(t)}{6},\,0,\,1\right)"),
    "6-25": ("latex", r"\kappa_i(t)=1-\operatorname{clip}\left(\frac{d_i^{\mathrm{center}}(t)}{0.7L_{\mathrm{map}}},\,0,\,1\right)"),
}


def set_run_font(
    run,
    *,
    latin: str = "Times New Roman",
    east_asia: str = "宋体",
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_line_style(
    paragraph,
    *,
    first_indent_cm: float = 0.0,
    space_before: float = 0,
    space_after: float = 0,
    line_pt: float = 20,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(first_indent_cm)
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)


def add_text_paragraph(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_indent_cm: float = 0.74,
    size: float = 10.5,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
    bold: bool = False,
    line_pt: float = 20,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_line_style(p, first_indent_cm=first_indent_cm, line_pt=line_pt)
    run = p.add_run(text)
    set_run_font(run, latin=latin, east_asia=east_asia, size=size, bold=bold)


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.74)
    fmt.first_line_indent = Cm(-0.74)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, latin="Times New Roman", east_asia="宋体", size=10.5)


def add_page_number(footer_paragraph) -> None:
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_run_font(run, latin="Times New Roman", east_asia="宋体", size=9)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(28)
    section.right_margin = Mm(22)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    run = p.add_run("基于集群的多任务协同态势感知平台")
    set_run_font(run, latin="Times New Roman", east_asia="黑体", size=20, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    run2 = p2.add_run("毕业论文初稿（汇总版）")
    set_run_font(run2, latin="Times New Roman", east_asia="黑体", size=18, bold=True)

    info = [
        "学生：贺小双",
        "专业：计算机相关方向",
        "说明：本稿依据当前项目实现内容汇总生成",
    ]
    for index, item in enumerate(info):
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info.paragraph_format.space_before = Pt(24 if index == 0 else 8)
        r = p_info.add_run(item)
        set_run_font(r, latin="Times New Roman", east_asia="宋体", size=12)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_run_font(r, latin="Times New Roman", east_asia="黑体", size=16, bold=True)
    elif level == 2:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_run_font(r, latin="Times New Roman", east_asia="黑体", size=13, bold=True)
    else:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        set_run_font(r, latin="Times New Roman", east_asia="黑体", size=11.5, bold=True)


def find_matching_paren(text: str, start: int) -> int:
    depth = 0
    for index in range(start, len(text)):
        char = text[index]
        if char == "(":
            depth += 1
        elif char == ")":
            depth -= 1
            if depth == 0:
                return index
    return -1


def replace_function_calls(text: str, name: str, builder) -> str:
    token = f"{name}("
    result: list[str] = []
    index = 0
    while index < len(text):
        pos = text.find(token, index)
        if pos == -1:
            result.append(text[index:])
            break
        result.append(text[index:pos])
        left = pos + len(name)
        right = find_matching_paren(text, left)
        if right == -1:
            result.append(text[pos:])
            break
        inner = text[left + 1 : right]
        result.append(builder(inner))
        index = right + 1
    return "".join(result)


def promote_simple_fractions(text: str) -> str:
    pattern = re.compile(r"([A-Za-z0-9_{}().\\]+)\s*/\s*([A-Za-z0-9_{}().\\]+)")
    previous = None
    current = text
    while previous != current:
        previous = current
        current = pattern.sub(r"\\frac{\1}{\2}", current)
    return current


def normalize_formula_to_latex(formula: str) -> str | None:
    lower = f" {formula.lower()} "
    if any(token in lower for token in (" otherwise", " if ", " and ", ";", " down", " up", " empty")):
        return None

    s = formula.strip()
    s = s.replace("<=", r"\leq ")
    s = s.replace(">=", r"\geq ")
    s = s.replace("!=", r"\neq ")
    s = s.replace(" cap ", r" \cap ")
    s = s.replace(" in ", r" \in ")
    s = s.replace(" and ", r" \text{ and } ")
    s = s.replace("|U_a(t)|", r"\left|U_a(t)\right|")
    s = s.replace("~", r"\tilde ")

    s = re.sub(r"\bDeltaE\b", r"\\Delta E", s)
    s = re.sub(r"\bDelta\b", r"\\Delta", s)

    token_patterns = [
        (r"\btheta(?=_)", r"\\theta"),
        (r"\bdelta(?=[_(])", r"\\delta"),
        (r"\bphi(?=[_(])", r"\\phi"),
        (r"\brho(?=_)", r"\\rho"),
        (r"\bsigma(?=[_(])", r"\\sigma"),
        (r"\bepsilon(?=_)", r"\\epsilon"),
        (r"\bbeta(?=[_(])", r"\\beta"),
        (r"\beta(?=[_(])", r"\\eta"),
        (r"\bchi(?=_)", r"\\chi"),
        (r"\btau(?=_)", r"\\tau"),
        (r"\bkappa(?=_)", r"\\kappa"),
        (r"\bOmega\b", r"\\Omega"),
        (r"\bPhi\b", r"\\Phi"),
        (r"\bGamma\b", r"\\Gamma"),
    ]
    for pattern, replacement in token_patterns:
        s = re.sub(pattern, replacement, s)

    s = replace_function_calls(s, "sqrt", lambda inner: rf"\sqrt{{{inner}}}")
    for fn in (
        "clip",
        "max",
        "min",
        "round",
        "sum",
        "lerp",
        "atan2",
        "dist_path",
        "dist_grid",
        "angle_diff",
        "exp",
    ):
        s = replace_function_calls(s, fn, lambda inner, fn=fn: rf"\operatorname{{{fn}}}\left({inner}\right)")

    s = re.sub(r"([A-Za-z\\]+)_([A-Za-z0-9]+)", r"\1_{\2}", s)
    s = re.sub(r"([A-Za-z\\}])\^\(([^)]+)\)", r"\1^{(\2)}", s)
    s = re.sub(r"([A-Za-z\\}\]])\^([A-Za-z0-9]+)", r"\1^{\2}", s)
    s = s.replace(" * ", r" \cdot ")
    s = promote_simple_fractions(s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def equation_image_path(number: str) -> Path:
    safe = number.replace("(", "").replace(")", "").replace("/", "_")
    return EQUATION_DIR / f"eq_{safe}.png"


def render_equation_image(number: str, formula: str) -> Path:
    EQUATION_DIR.mkdir(parents=True, exist_ok=True)
    target = equation_image_path(number or "eq")

    number_key = number.strip("()")
    initial_mode, initial_expr = EQUATION_OVERRIDES.get(
        number_key, ("latex", normalize_formula_to_latex(formula) or formula)
    )

    def draw(mode: str, expr: str) -> None:
        text_length = len(expr.replace("\n", " "))
        if mode == "piecewise":
            height = 1.35
            width = 9.2
        else:
            height = 0.95 if "\n" not in expr else 1.25
            width = max(7.6, min(10.4, 6.8 + text_length * 0.032))

        fig = plt.figure(figsize=(width, height), dpi=260)
        fig.patch.set_alpha(0)
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")

        if mode == "piecewise":
            if expr == "chi":
                ax.text(0.39, 0.50, r"$\chi_O(x,y)=$", ha="right", va="center", fontsize=18.5, color="black")
                ax.text(0.425, 0.50, "{", ha="center", va="center", fontsize=34, color="black")
                ax.text(0.47, 0.63, "1,  (x,y)∈O", ha="left", va="center", fontsize=15.8, color="black")
                ax.text(0.47, 0.37, "0,  (x,y)∉O", ha="left", va="center", fontsize=15.8, color="black")
            else:
                ax.text(0.39, 0.50, r"$V_{ij}(t)=$", ha="right", va="center", fontsize=18.5, color="black")
                ax.text(0.425, 0.50, "{", ha="center", va="center", fontsize=34, color="black")
                ax.text(0.47, 0.63, "1,  line(u_i,t_j)∩O=∅", ha="left", va="center", fontsize=15.3, color="black")
                ax.text(0.47, 0.37, "0,  otherwise", ha="left", va="center", fontsize=15.3, color="black")
        elif mode == "text":
            font_size = 15.5 if "\n" not in expr else 14.0
        else:
            font_size = 17.5
            if text_length > 70:
                font_size = 16.0
            if text_length > 100:
                font_size = 14.3
            if text_length > 130:
                font_size = 13.0

        if mode == "latex":
            rendered = expr if expr.startswith("$") else f"${expr}$"
            ax.text(0.455, 0.50, rendered, ha="center", va="center", fontsize=font_size, color="black")
        elif mode == "text":
            plain_expr = (
                expr.replace("<=", "≤")
                .replace(">=", "≥")
                .replace("!=", "≠")
                .replace(" cap ", " ∩ ")
                .replace(" in ", " ∈ ")
            )
            ax.text(
                0.455,
                0.50,
                plain_expr,
                ha="center",
                va="center",
                fontsize=font_size,
                color="black",
                linespacing=1.12,
            )

        ax.text(0.982, 0.50, f"({number_key})", ha="right", va="center", fontsize=12.5, color="black")
        fig.savefig(target, dpi=260, bbox_inches="tight", pad_inches=0.08, transparent=True)
        plt.close(fig)

    try:
        draw(initial_mode, initial_expr)
    except Exception:
        draw("text", formula)
    return target


def add_equation_block(doc: Document, raw_formula: str) -> None:
    text = raw_formula.strip()
    match = EQUATION_RE.match(text)
    number = ""
    formula = text
    if match:
        number = f"({match.group(1)})"
        formula = match.group(2).strip()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(0.0)
    fmt.space_before = Pt(10)
    fmt.space_after = Pt(12)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(render_equation_image(number, formula)), width=Cm(15.0))


def parse_markdown_to_docx(doc: Document, content: str) -> None:
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = " ".join(line.strip() for line in paragraph_lines).strip()
        if text:
            if text.startswith("[") and "]" in text:
                add_reference_paragraph(doc, text)
            else:
                add_text_paragraph(doc, text)
        paragraph_lines = []

    for raw in content.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped == "---":
            flush_paragraph()
            doc.add_page_break()
            continue

        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            flush_paragraph()
            add_equation_block(doc, stripped[2:-2].strip())
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            add_heading(doc, stripped[2:], 1)
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            add_heading(doc, stripped[3:], 2)
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            add_heading(doc, stripped[4:], 3)
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_text_paragraph(doc, "• " + stripped[2:], first_indent_cm=0.0, line_pt=18)
            continue

        paragraph_lines.append(stripped)

    flush_paragraph()


def generate_docx() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    EQUATION_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    content = SOURCE.read_text(encoding="utf-8")
    parse_markdown_to_docx(doc, content)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = generate_docx()
    print(path)
