# -*- coding: utf-8 -*-
from __future__ import unicode_literals

from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = r"D:\biushe\tmp\docs\thesis_v2.docx"
TARGET = r"D:\biushe\output\doc\毕业论文初稿_汇总版_公式终版_v2_插入符号表.docx"


ROWS = [
    (u"场景建模", u"𝓔(t)", u"第 t 步的整体实验场景", u"包含任务区域、障碍物、无人机、目标和天气"),
    (u"场景建模", u"Ω", u"任务区域", u"平台中的任务空间"),
    (u"场景建模", u"O", u"障碍物集合", u"包括建筑物和不可通行区域"),
    (u"场景建模", u"U(t)", u"第 t 步的无人机状态集合", u"所有无人机节点"),
    (u"场景建模", u"T(t)", u"第 t 步的目标状态集合", u"所有动态目标"),
    (u"场景建模", u"W", u"天气状态", u"晴空、薄雾、降雨、雷暴"),
    (u"场景建模", u"χ_O(x,y)", u"障碍物占用函数", u"判断位置 (x,y) 是否被障碍物占据"),
    (u"场景建模", u"V_ij(t)", u"视线可达函数", u"判断无人机 u_i 与目标 t_j 之间是否无遮挡"),
    (u"无人机模型", u"u_i(t)", u"第 i 架无人机在时刻 t 的状态", u"包含位置、朝向、能量和偏置等"),
    (u"无人机模型", u"(x_i(t), y_i(t), z_i(t))", u"无人机空间位置", u"分别表示平面位置和飞行高度"),
    (u"无人机模型", u"θ_i(t)", u"无人机朝向角", u"用于视场判定与姿态显示"),
    (u"无人机模型", u"E_i(t)", u"无人机剩余能量", u"随仿真推进不断变化"),
    (u"无人机模型", u"b_i^v", u"无人机速度偏置", u"描述不同节点飞行能力差异"),
    (u"无人机模型", u"b_i^s", u"无人机感知偏置", u"描述不同节点感知能力差异"),
    (u"无人机模型", u"b_i^c", u"无人机通信偏置", u"描述不同节点链路能力差异"),
    (u"无人机模型", u"F_i(t)", u"无人机故障状态", u"表示节点是否失效"),
    (u"无人机模型", u"v_d", u"无人机基准速度", u"对应系统配置中的无人机速度参数"),
    (u"无人机模型", u"L_i^max(t)", u"当前步无人机最大可移动距离", u"受天气和速度偏置共同影响"),
    (u"无人机模型", u"ΔE_i(t)", u"当前步无人机总能耗", u"包括移动、通信、待机和高度调整能耗"),
    (u"无人机模型", u"A_i(t)", u"无人机可用状态函数", u"判断该节点当前是否可参与任务"),
    (u"无人机模型", u"E_min", u"最小可用能量阈值", u"低于该值时节点退出任务"),
    (u"目标模型", u"t_j(t)", u"第 j 个目标在时刻 t 的状态", u"包含位置、速度、优先级和不确定性"),
    (u"目标模型", u"(x_j(t), y_j(t))", u"目标位置", u"目标在任务平面中的当前坐标"),
    (u"目标模型", u"(v_xj(t), v_yj(t))", u"目标速度分量", u"描述目标运动状态"),
    (u"目标模型", u"p_j", u"目标优先级", u"表示目标重要程度"),
    (u"目标模型", u"q_j(t)", u"目标不确定性", u"表示当前系统对目标掌握程度"),
    (u"目标模型", u"v_j(t)", u"目标速度模长", u"由速度分量计算得到"),
    (u"目标模型", u"v_t", u"目标基准速度", u"对应系统配置中的目标速度参数"),
    (u"路径规划", u"G(r,c)", u"栅格地图占用状态", u"判断网格是否可通行"),
    (u"路径规划", u"M_j(r,c)", u"目标 t_j 的距离图", u"表示任意网格到目标的最短路径距离"),
    (u"路径规划", u"(r_i,c_i)", u"无人机所在网格坐标", u"由当前位置映射得到"),
    (u"路径规划", u"(r_j,c_j)", u"目标所在网格坐标", u"由目标位置映射得到"),
    (u"路径规划", u"l_cell", u"单个网格边长", u"栅格地图的物理尺度"),
    (u"路径规划", u"D_ij", u"无人机到目标的绕障路径代价", u"用于调度评分与运动更新"),
    (u"路径规划", u"𝒩_4(r_i,c_i)", u"当前网格的四邻域可通行单元集合", u"用于寻找下一跳路径点"),
    (u"通信约束", u"ρ_i(t)", u"邻居密度得分", u"表示通信半径内可用邻居数量的归一化结果"),
    (u"通信约束", u"κ_i(t)", u"中心位置得分", u"反映节点相对地图中心的位置优势"),
    (u"通信约束", u"L_i(t)", u"链路质量", u"受邻居密度、中心位置、偏置和天气影响"),
    (u"通信约束", u"p_0", u"基础丢包率", u"系统配置中的基础通信丢包率"),
    (u"通信约束", u"Δp(W)", u"天气附加丢包量", u"不同天气对丢包率的额外影响"),
    (u"通信约束", u"p_i^loss(t)", u"有效丢包率", u"当前链路条件下的实际丢包概率"),
    (u"通信约束", u"d_0", u"基础延迟步数", u"系统配置中的通信延迟"),
    (u"通信约束", u"Δd(W)", u"天气附加延迟", u"恶劣天气引入的额外延迟步数"),
    (u"通信约束", u"τ_i^max(t)", u"当前通信延迟预算", u"当前链路可产生的最大延迟步数"),
    (u"天气扰动", u"R_s", u"基础感知半径", u"系统配置中的感知范围参数"),
    (u"天气扰动", u"R_s^(W)", u"天气作用下的有效感知半径", u"感知半径经天气缩放后的结果"),
    (u"天气扰动", u"Φ_s^(W)", u"天气作用下的有效视场角", u"视场角经天气缩放后的结果"),
    (u"天气扰动", u"σ_s", u"基础观测噪声", u"系统配置中的传感器噪声参数"),
    (u"天气扰动", u"σ_s^(W)", u"天气作用下的有效观测噪声", u"基础噪声经天气放大后的结果"),
    (u"天气扰动", u"p_f", u"基础故障率", u"系统配置中的基础故障概率"),
    (u"天气扰动", u"p_f^(W)", u"天气作用下的故障率", u"考虑天气附加量后的节点故障概率"),
    (u"天气扰动", u"Δf(W)", u"天气附加故障量", u"恶劣天气带来的附加故障概率"),
    (u"感知判定", u"d_ij(t)", u"无人机与目标的欧氏距离", u"目标感知的首要几何约束"),
    (u"感知判定", u"c_ij(t)", u"覆盖度", u"反映目标相对感知中心区域的接近程度"),
    (u"感知判定", u"P_ij^det(t)", u"探测概率", u"当前条件下目标被成功探测的概率"),
    (u"感知判定", u"z_ij(t)", u"无人机对目标的观测值", u"为带噪观测结果"),
    (u"感知判定", u"ε_x,ij(t), ε_y,ij(t)", u"观测噪声项", u"表示观测的随机误差"),
    (u"感知判定", u"σ_ij(t)", u"当前观测噪声标准差", u"与天气、距离和感知偏置有关"),
    (u"感知判定", u"w_ij^0(t)", u"基础观测权重", u"融合前观测结果的初始权重"),
    (u"RADS 调度", u"P_j(t)", u"目标优先级归一化项", u"在 6.7.2 中使用"),
    (u"RADS 调度", u"Q_j(t)", u"目标不确定性归一化项", u"在 6.7.2 中使用"),
    (u"RADS 调度", u"M_j(t)", u"目标速度强度项", u"在 6.7.2 中使用，注意与距离图 M_j(r,c) 含义不同"),
    (u"RADS 调度", u"ξ_j(t)", u"目标需求扰动项", u"用于增加小幅动态变化"),
    (u"RADS 调度", u"R_j(t)", u"目标紧急度", u"反映目标当前的综合需求水平"),
    (u"RADS 调度", u"N_max", u"单目标需求上限", u"当前步单个目标允许的最大派机规模"),
    (u"RADS 调度", u"N_j^dem(t)", u"目标需求规模", u"当前步目标需要的无人机数量"),
    (u"RADS 调度", u"C_ij(t)", u"路径覆盖度项", u"在效用函数中反映路径可达性"),
    (u"RADS 调度", u"E_i'(t)", u"能量归一化项", u"在效用函数中表示节点剩余能量水平"),
    (u"RADS 调度", u"S_ij(t)", u"感知就绪项", u"反映节点当前是否接近有效感知状态"),
    (u"RADS 调度", u"B_ij^los(t)", u"视线奖励项", u"反映无遮挡带来的额外加分"),
    (u"RADS 调度", u"P_ij^travel(t)", u"路径惩罚项", u"反映路径过长带来的代价"),
    (u"RADS 调度", u"Γ_j(t)", u"冗余度项", u"反映目标当前已获支持程度"),
    (u"RADS 调度", u"U_ij(t)", u"无人机对目标的综合效用值", u"RADS 选取节点的核心依据"),
    (u"RADS 调度", u"Ē(t)", u"当前可用无人机平均能量", u"用于动态预算估计"),
    (u"RADS 调度", u"L̄(t)", u"当前可用无人机平均链路质量", u"用于动态预算估计"),
    (u"RADS 调度", u"R_Σ(t)", u"全部目标的总紧急度", u"用于估计总体任务压力"),
    (u"RADS 调度", u"H(t)", u"热点目标数量", u"表示高需求或高紧急度目标数"),
    (u"RADS 调度", u"P_recent(t)", u"近期定位达标率", u"用于反映近期感知表现"),
    (u"RADS 调度", u"B*(t)", u"期望派机预算", u"动态预算的中间估计值"),
    (u"RADS 调度", u"B(t)", u"最终动态预算", u"当前步系统应派出的无人机总数"),
    (u"随机派遣", u"B_rand(t)", u"随机派遣预算", u"当前实现中与 RADS 共用动态预算"),
    (u"全量派遣", u"U_a(t)", u"当前可用无人机集合", u"未故障且能量高于阈值的节点集合"),
    (u"全量派遣", u"S^full(t)", u"全量派遣子群", u"当前步全部可用无人机组成的执行集合"),
    (u"全量派遣", u"B_full(t)", u"全量派遣规模", u"当前步可用无人机总数"),
    (u"全量派遣", u"Λ_j(t)", u"目标支持比", u"表示当前目标已获支持程度"),
    (u"全量派遣", u"B_j^sup(t)", u"目标支持修正项", u"用于调节全量派遣时的目标选择偏好"),
    (u"全量派遣", u"S_ij^full(t)", u"全量派遣评分", u"由综合效用和支持修正项共同构成"),
    (u"融合判定", u"x_c(t), y_c(t)", u"粗加权中心坐标", u"融合前的粗估计中心"),
    (u"融合判定", u"r_n(t)", u"第 n 条观测的残差", u"观测相对粗中心的偏离程度"),
    (u"融合判定", u"R_0(t)", u"基准残差半径", u"用于稳健权重修正"),
    (u"融合判定", u"w_n'(t)", u"稳健权重", u"融合时对异常观测降权后的权重"),
    (u"融合判定", u"N_f", u"最终参与融合的观测数", u"经过筛选后保留的观测数量"),
    (u"融合判定", u"(x̂_j(t), ŷ_j(t))", u"融合估计位置", u"目标最终估计结果"),
    (u"融合判定", u"e_j(t)", u"融合误差", u"融合结果与目标真实位置之间的偏差"),
    (u"融合判定", u"A_j(t)", u"信息时效", u"当前融合观测相对当前时刻的平均滞后"),
    (u"融合判定", u"s_j(t)", u"观测离散程度", u"用于衡量多观测之间的分散程度"),
    (u"融合判定", u"g_j(t)", u"一致性指标", u"用于衡量多观测是否集中一致"),
    (u"严格判定", u"N_j^conf(t)", u"目标确认门限", u"当前步目标所需的最少有效观测数"),
    (u"严格判定", u"C_j(t)", u"确认条件", u"判断观测数是否达到确认门限"),
    (u"严格判定", u"P_j^pass(t)", u"定位达标条件", u"建议用该记号避免与 6.7.2 中 P_j(t) 重复"),
    (u"严格判定", u"G_j(t)", u"一致性达标条件", u"判断一致性是否达到平台阈值"),
    (u"严格判定", u"S_j(t)", u"严格成功条件", u"由确认、误差和一致性共同决定"),
    (u"总体指标", u"N_success", u"严格成功总次数", u"用于计算严格成功率"),
    (u"总体指标", u"N_attempt", u"总尝试次数", u"目标-时间步级别的尝试总数"),
    (u"总体指标", u"N_pass", u"定位达标总次数", u"用于计算定位达标率"),
    (u"总体指标", u"N_confirm", u"确认成功总次数", u"用于计算确认率"),
    (u"总体指标", u"SR", u"累计严格成功率", u"系统核心评价指标之一"),
    (u"总体指标", u"PR", u"累计定位达标率", u"系统核心评价指标之一"),
    (u"总体指标", u"CR", u"累计确认率", u"系统核心评价指标之一"),
    (u"复杂度分析", u"k", u"目标数量", u"当前场景中的目标数"),
    (u"复杂度分析", u"m", u"无人机数量", u"当前场景中的无人机数"),
    (u"复杂度分析", u"|G|", u"地图规模", u"栅格地图的规模"),
    (u"复杂度分析", u"N_z", u"当前步总有效观测数", u"用于估计融合计算开销"),
    (u"复杂度分析", u"𝒪(·)", u"时间复杂度记号", u"用于表示算法复杂度数量级"),
]


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = u""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    run.font.size = Pt(size)


def set_row_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def style_table(table):
    table.style = "Table Grid"
    table.autofit = False
    widths = [Pt(60), Pt(100), Pt(170), Pt(150)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def insert_before_paragraph(doc, target_para, title_text, rows):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title_text)
    title_run.bold = True
    title_run.font.name = u"黑体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")
    title_run.font.size = Pt(12)
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note_p.add_run(u"注：本表对第六章中主要符号进行统一说明，个别重复符号已结合上下文作区分处理。")
    note_run.font.name = u"宋体"
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    note_run.font.size = Pt(10.5)
    note_p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    hdr = table.rows[0].cells
    headers = [u"类别", u"符号", u"含义", u"备注"]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    set_row_header(table.rows[0])

    for category, symbol, meaning, note in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], category)
        set_cell_text(cells[1], symbol)
        set_cell_text(cells[2], meaning)
        set_cell_text(cells[3], note)

    # move created elements before target paragraph
    target = target_para._p
    title_el = deepcopy(title_p._p)
    note_el = deepcopy(note_p._p)
    table_el = deepcopy(table._tbl)
    target.addprevious(table_el)
    target.addprevious(note_el)
    target.addprevious(title_el)

    # remove temporary end elements
    title_p._element.getparent().remove(title_p._element)
    note_p._element.getparent().remove(note_p._element)
    table._element.getparent().remove(table._element)


def main():
    doc = Document(SOURCE)
    target_para = None
    for p in doc.paragraphs:
        if p.text.strip() == u"6.12 本章小结":
            target_para = p
            break
    if target_para is None:
        target_para = doc.add_paragraph()

    insert_before_paragraph(doc, target_para, u"表6-1 第六章主要符号说明表", ROWS)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
